#
# Imports :
#
from datetime import datetime
from django.utils import timezone
from django.http import HttpResponse
from django.views.generic import View
from varcaldev.utils.cEmpresa import *
from varcaldev.models.mpaquetes import Bulto
from varcaldev.models.mguia import Guia
from django.urls import reverse
from django.contrib import messages
from django.contrib.messages.views import SuccessMessageMixin
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
import os
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
#
#  name: class GenerateExcel
#
class GenerateEtiquetas(SuccessMessageMixin, View):
    def get(self, request, *args, **kwargs):
        id = str
        id = kwargs["pk"]
        query_m = Guia.objects.filter(id__exact=id)
        recolecta = query_m[0].recolecta
        model = Bulto
        results = Bulto.objects.filter(recolecta__exact=recolecta)
        total = results.count()
        document = Document(os.path.join(BASE_DIR, 'excel_views/labels-base.docx'))
        #document = Document()
        style = document.styles['Normal']
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(0.2)
            section.bottom_margin = Cm(0.3)
            section.left_margin = Cm(0.3)
            section.right_margin = Cm(0.3)
        font = style.font
        font.name = 'Arial' 
        font.size = Pt(9)
        if total <= 50:
            emp = EmpresaDatos()
            i = 0
            for result in results:
                table = document.add_table(rows=0, cols=2)
                #
                rcA = table.add_row().cells
                p_rcAl = rcA[0].add_paragraph(results[i].destino)
                p_rcAl.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pr_rcAl = rcA[0].paragraphs
                for paragraph in pr_rcAl:
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Arial'
                        font.size = Pt(10)
                p_rcAr = rcA[1].add_paragraph('B: ' + str(i+1) + '/' + str(total))
                p_rcAr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                pr_rcAr = rcA[1].paragraphs
                for paragraph in pr_rcAr:
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Arial'
                        font.size = Pt(10)
                #
                lB = document.add_paragraph(str(results[i].ruta))
                #
                lC = document.add_paragraph(str(results[i].destinatario))
                #
                lD = document.add_paragraph('Guía: ' + str(results[i].guias) + '  / Recolecta: ' + str(results[i].recolecta))
                #
                lE = document.add_paragraph(str(results[i].origen))
                lE.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                #
                lF = document.add_paragraph()
                lF.alignment = WD_ALIGN_PARAGRAPH.CENTER
                runF = lF.add_run('*' + results[i].id_bulto + '*')
                runF.font.name = 'C39HrP24DhTt'
                runF.font.size = Pt(28)
                #
                lG = document.add_paragraph(emp.nombre + ' - Rif: ' + emp.rif)
                lG.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                i = i + 1
                document.add_page_break()
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename={date}-Etiquetas.docx'.format(
                date=datetime.now().strftime('%d-%m-%Y'),
            )
            document.save(response)
            return response
        else:
            emp = EmpresaDatos()
            table = document.add_table(rows=0, cols=2)
            rcA = table.add_row().cells
            p_rcAl = rcA[0].add_paragraph(results[0].destino)
            p_rcAl.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pr_rcAl = rcA[0].paragraphs
            for paragraph in pr_rcAl:
                for run in paragraph.runs:
                    font = run.font
                    font.name = 'Arial'
                    font.size = Pt(10)
            p_rcAr = rcA[1].add_paragraph('B: ' + str(total))
            p_rcAr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pr_rcAr = rcA[1].paragraphs
            for paragraph in pr_rcAr:
                for run in paragraph.runs:
                    font = run.font
                    font.name = 'Arial'
                    font.size = Pt(10)
            #
            lB = document.add_paragraph(str(results[0].ruta))
            #
            lC = document.add_paragraph(str(results[0].destinatario))
            #
            lD = document.add_paragraph('Guía: ' + str(results[0].guias) + '  / Recolecta: ' + str(results[0].recolecta))
            #
            lE = document.add_paragraph(str(results[0].origen))
            lE.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            #
            lF = document.add_paragraph()
            lF.alignment = WD_ALIGN_PARAGRAPH.CENTER
            runF = lF.add_run('*' + results[0].id_bulto + '*')
            runF.font.name = 'C39HrP24DhTt'
            runF.font.size = Pt(28)
            #
            lG = document.add_paragraph(emp.nombre + ' - Rif: ' + emp.rif)
            lG.alignment = WD_ALIGN_PARAGRAPH.CENTER
            #
            document.add_page_break()
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename={date}-Etiquetas.docx'.format(
                date=datetime.now().strftime('%d-%m-%Y'),
            )
            document.save(response)
            return response